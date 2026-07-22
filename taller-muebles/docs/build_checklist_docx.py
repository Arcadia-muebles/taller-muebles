from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "checklist-aceptacion-segundo-pago-rodrigo.md"
OUTPUT = ROOT / "checklist-aceptacion-segundo-pago-rodrigo.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
GREEN = "2F6B4F"
RED = "A33A3A"
TEXT = "252525"
MUTED = "666666"
BORDER = "C9D3DE"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=BORDER, size=5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_inline(paragraph, text: str, default_bold=False, default_color=TEXT) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        bold = default_bold
        mono = False
        content = part
        if part.startswith("**") and part.endswith("**"):
            content = part[2:-2]
            bold = True
        elif part.startswith("`") and part.endswith("`"):
            content = part[1:-1]
            mono = True
        run = paragraph.add_run(content)
        run.font.name = "Consolas" if mono else "Calibri"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), run.font.name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), run.font.name)
        run.font.size = Pt(9.5 if mono else 11)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(default_color)


def add_check(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Checklist")
    p.paragraph_format.keep_together = True
    box = p.add_run("☐")
    box.font.name = "Segoe UI Symbol"
    box._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Segoe UI Symbol")
    box._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Segoe UI Symbol")
    box.font.size = Pt(12)
    box.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run("  ")
    add_inline(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    add_inline(p, text)


def add_response_line(doc: Document, label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    if label:
        run = p.add_run(label + " ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    line = p.add_run("_" * (72 if not label else 48))
    line.font.color.rgb = RGBColor.from_string("9AA7B4")


def add_callout(doc: Document, label: str, body: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], 120)
    set_table_borders(table, BORDER, 5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(label.upper())
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    add_inline(body_p, body)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    checklist = doc.styles.add_style("Checklist", 1)
    checklist.font.name = "Calibri"
    checklist.font.size = Pt(11)
    checklist.paragraph_format.left_indent = Inches(0.375)
    checklist.paragraph_format.first_line_indent = Inches(-0.375)
    checklist.paragraph_format.space_after = Pt(4)
    checklist.paragraph_format.line_spacing = 1.16


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("ARCADIA  ·  VALIDACIÓN FUNCIONAL")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Checklist de aceptación")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Revisión de flujos para autorizar el segundo pago")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=5, cols=2)
    set_table_geometry(table, [2400, 6960], 120)
    set_table_borders(table, BORDER, 5)
    rows = [
        ("Proyecto", "ARCADIA"),
        ("Revisor", "Rodrigo"),
        ("Fecha de revisión", ""),
        ("Ambiente / URL", ""),
        ("Versión o commit", ""),
    ]
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        set_cell_shading(left, LIGHT_BLUE)
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.add_run(value if value else "____________________________________________")

    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    add_callout(
        doc,
        "Condición de aprobación",
        "Todas las pruebas BLOQUEANTES deben quedar en OK o ser aceptadas expresamente por Rodrigo como N/A. Debe completarse al menos un pedido ficticio de punta a punta y comprobarse la persistencia después de cerrar sesión.",
    )


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", cell) for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = [500, 1700, 2780, 1180, 1200, 1200, 800]
    if len(rows[0]) != 7:
        widths = [9360 // len(rows[0])] * len(rows[0])
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths, 120)
    set_table_borders(table, BORDER, 5)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(value)
            r.bold = row_idx == 0
            r.font.name = "Calibri"
            r.font.size = Pt(8.5 if row_idx else 8)
            r.font.color.rgb = RGBColor.from_string(DARK_BLUE if row_idx == 0 else TEXT)


def build() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "ARCADIA  |  Checklist de aceptación funcional"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.runs[0]
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(section.footer.paragraphs[0])

    add_title_block(doc)

    idx = 0
    # Skip the Markdown title and metadata already represented in the title block.
    while idx < len(lines) and not lines[idx].startswith("## Objetivo"):
        idx += 1

    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            idx += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("## "):
            heading = stripped[3:]
            if heading == "Objetivo":
                p = doc.add_paragraph(heading, style="Heading 1")
            elif re.match(r"^\d+\.", heading):
                p = doc.add_paragraph(heading, style="Heading 1")
            else:
                p = doc.add_paragraph(heading, style="Heading 2")
            set_keep_with_next(p)
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 3")
        elif stripped.startswith("- [ ] "):
            add_check(doc, stripped[6:])
        elif stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
        elif re.fullmatch(r"_+", stripped):
            add_response_line(doc)
        elif stripped.endswith(":**") and any(label in stripped for label in ("Evidencia", "Pedido", "Áreas", "Dispositivo")):
            label = stripped.removeprefix("**").removesuffix("**").removesuffix(":")
            add_response_line(doc, label + ":")
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        idx += 1

    props = doc.core_properties
    props.title = "Checklist de aceptación funcional — segundo pago"
    props.subject = "Validación de flujos operacionales de ARCADIA"
    props.author = "ARCADIA"
    props.keywords = "ARCADIA, checklist, aceptación, QA, segundo pago"
    props.comments = "Documento preparado para revisión funcional de Rodrigo."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
